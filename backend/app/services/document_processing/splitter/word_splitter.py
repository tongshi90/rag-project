"""
Word 文档拆分模块

负责将 Word (.docx) 文档解析并拆分为文本 chunks。
复用 PDF 分片的标题识别和分片逻辑。

【主入口函数】
- split_word_to_chunks(word_path)  # 从 Word 文件路径直接生成分片
"""
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple

try:
    from docx import Document
except ImportError:
    raise ImportError("请先安装 python-docx: pip install python-docx")

# 导入 PDF 分片中的复用函数
from .text_splitter import (
    refine_title_patterns,
    split_chunks,
    post_process_chunks,
    extract_keywords,
    RAW_TITLE_PATTERNS
)


def extract_paragraphs_from_word(word_path: str) -> List[Dict]:
    """
    从 Word 文档中提取所有段落，转换为与 PDF 兼容的格式

    Args:
        word_path: Word 文件路径

    Returns:
        List[Dict], 每个包含 {"text": str, "page_number": int}
    """
    doc = Document(word_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检查段落样式，如果是标题样式，可以添加标记
        # 确保样式名称是字符串类型
        try:
            style_name = str(para.style.name) if (para.style and hasattr(para.style, 'name')) else ""
        except (AttributeError, TypeError):
            style_name = ""

        lines.append({
            "text": text,
            "page_number": 1,  # Word 文档没有页码概念，统一使用 1（注意：key 是 page_number 不是 page）
            "style": style_name,  # 保留样式信息供后续参考
        })

    # 处理表格中的文本（可选）
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                lines.append({
                    "text": " | ".join(row_text),  # 表格行用 | 分隔
                    "page_number": 1,
                    "style": "Table",
                })

    return lines


def refine_title_patterns_for_word(lines: List[Dict]) -> List:
    """
    Word 文档的标题规则优化

    优先利用 Word 内置的标题样式，如果样式不足则使用文本模式匹配

    Args:
        lines: 文本行列表

    Returns:
        标题正则规则列表
    """
    # 首先检查 Word 内置标题样式的使用情况
    heading_styles = set()
    for line in lines:
        style = line.get("style", "")
        if "Heading" in style or "标题" in style:
            heading_styles.add(style)

    # 如果 Word 文档使用了内置标题样式，可以基于样式构建规则
    # 这里为了简化，我们仍然使用文本模式匹配（兼容性更好）
    # 但可以优先检查带有标题样式的段落

    # 使用与 PDF 相同的标题检测逻辑
    return refine_title_patterns(lines)


def split_word_to_chunks(word_path: str, file_id: str = None, kb_id: str = None) -> tuple[List[Dict], List]:
    """
    从 Word 文件直接生成 chunks

    处理流程：
    1. 提取段落（extract_paragraphs_from_word）
    2. 标题规则优化（refine_title_patterns_for_word）
    3. 按标题分片（split_chunks）
    4. 后处理（post_process_chunks）

    Args:
        word_path: Word 文件路径
        file_id: 文档 ID
        kb_id: 知识库 ID（可选）

    Returns:
        (chunks, title_patterns): 分片列表和标题正则规则
    """
    # 第一步：提取段落
    lines = extract_paragraphs_from_word(word_path)

    if not lines:
        return [], []

    # 第二步：标题规则优化
    title_patterns = refine_title_patterns_for_word(lines)

    # 第三步：按标题分片
    chunks = split_chunks(lines, title_patterns)

    # 第四步：后处理
    chunks = post_process_chunks(chunks, file_id, kb_id)

    return chunks, title_patterns


def is_word_file(file_path: str) -> bool:
    """
    判断是否为 Word 文档

    Args:
        file_path: 文件路径

    Returns:
        是否为 .doc 或 .docx 文件
    """
    ext = Path(file_path).suffix.lower()
    return ext in ['.doc', '.docx']


if __name__ == "__main__":
    # 测试代码
    word_path = r"F:\rag_project\backend\upload\sample.docx"
    file_id = "1770601282.390849"
    text_chunks, title_patterns = split_word_to_chunks(word_path, file_id)

    # 打印分片信息
    print(f"\n========== Word 文档分片结果 ==========")
    print(f"分片数量: {len(text_chunks)}")
    print(f"标题规则: {title_patterns}")
    print(f"{'=' * 40}")

    import json

    for i, chunk in enumerate(text_chunks, 1):
        print(f"\n--- 分片 {i} ---")
        print(json.dumps(chunk, ensure_ascii=False, indent=2))

    print(f"\n{'=' * 40}\n")
